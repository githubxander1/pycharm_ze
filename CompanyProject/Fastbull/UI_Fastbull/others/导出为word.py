import os
import pdfkit
from jinja2 import Environment, FileSystemLoader
from docx import Document

def read_files_in_dir(directory):
    """ 递归读取目录下所有文件的内容 """
    result = []
    for root, dirs, files in os.walk(directory):
        level = root.replace(directory, '').count(os.sep)
        indent = ' ' * 4 * (level)
        # 将目录名和空字符串作为一个元组添加
        result.append((f"{indent}{'-' * level} {os.path.basename(root)}/", ""))
        sub_indent = ' ' * 4 * (level + 1)
        for f in files:
            filepath = os.path.join(root, f)
            try:
                with open(filepath, 'r', encoding='utf-8') as file:
                    content = file.read()
                    # 将文件名和内容作为一个元组添加
                    result.append((f"{sub_indent}{'-' * (level+1)} {f}:", content))
            except Exception as e:
                # 当文件无法读取时，将文件名和错误消息作为一个元组添加
                result.append((f"{sub_indent}{'-' * (level+1)} {f}: [无法读取文件]", ""))
    return result


def write_to_word(doc, items):
    """ 将文件名和内容写入Word文档 """
    for item in items:
        if item[0]:
            doc.add_paragraph(item[0])
        if item[1]:
            doc.add_paragraph(item[1])

if __name__ == "__main__":
    directory = r'D:\1test\PycharmProject\CompanyProject\Fastbull\Api_fastbull'
    items = read_files_in_dir(directory)

    # 创建Word文档
    doc = Document()

    # 写入内容
    write_to_word(doc, items)

    # 保存Word文档
    doc.save('output.docx')

    print("Word文档已创建.")
# def generate_pdf(html_content, output_file):
#     """ 将HTML内容转换为PDF """
#     options = {
#         'page-size': 'Letter',
#         'margin-top': '0.75in',
#         'margin-right': '0.75in',
#         'margin-bottom': '0.75in',
#         'margin-left': '0.75in',
#         'encoding': "UTF-8",
#     }
#     pdfkit.from_string(html_content, output_file, options=options)
#
# def create_html_template(content):
#     """ 创建HTML模板并填充内容 """
#     env = Environment(loader=FileSystemLoader('.'))
#     template = env.get_template('template.html')
#     html_content = template.render(content=content)
#     return html_content
#
# if __name__ == "__main__":
#     directory = r'D:\1test\PycharmProject\CompanyProject\Fastbull\Api_fastbull'
#     content = read_files_in_dir(directory)
#
#     # 创建HTML模板
#     html_content = create_html_template(content)
#
#     # 生成PDF
#     output_file = 'output.pdf'
#     generate_pdf(html_content, output_file)
#
#     print(f"PDF has been created at {output_file}")
