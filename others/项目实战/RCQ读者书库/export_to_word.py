import os
from docx import Document


def export_to_word(folder_path, output_doc_path):
    # 创建一个新的Word文档
    doc = Document()

    # 遍历文件夹中的所有文件
    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)

        # 检查是否是文件
        if os.path.isfile(file_path):
            # 读取文件内容
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()

            # 将文件内容添加到Word文档中
            doc.add_paragraph(f'文件名: {filename}')
            doc.add_paragraph(content)
            doc.add_paragraph('-' * 40)  # 添加分隔线

    # 保存Word文档
    doc.save(output_doc_path)
    print(f'文档已保存到 {output_doc_path}')


# 指定文件夹路径和输出Word文档的路径
folder_path = r'/others/项目实战/status_bar_notification'  # 替换为你的文件夹路径
output_doc_path = 'test1/output.docx'  # 你希望保存的Word文档路径

# 调用函数
export_to_word(folder_path, output_doc_path)