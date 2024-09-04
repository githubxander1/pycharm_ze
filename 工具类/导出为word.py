import os
from docx import Document  # 正确的导入方式

def read_files_in_dir(directory):
    """ 递归读取目录下所有文件的内容 """
    result = []
    for root, dirs, files in os.walk(directory):
        level = root.replace(directory, '').count(os.sep)
        indent = ' ' * 4 * (level)
        result.append((f"{indent}{'-' * level} {os.path.basename(root)}/", ""))
        sub_indent = ' ' * 4 * (level + 1)
        for f in files:
            filepath = os.path.join(root, f)
            try:
                with open(filepath, 'r', encoding='utf-8') as file:
                    content = file.read()
                    result.append((f"{sub_indent}{'-' * (level+1)} {f}:", content))
            except Exception as e:
                result.append((f"{sub_indent}{'-' * (level+1)} {f}: [无法读取文件]", ""))
    return result

def write_to_word(doc, items):
    """ 将文件名和内容写入Word文档 """
    for item in items:
        if item[0]:
            doc.add_paragraph(item[0])
        if item[1]:
            doc.add_paragraph(item[1])

if __name__ == "__main__":
    directory = r'D:\1test\PycharmProject\others\项目实战\局域网文件互传'
    items = read_files_in_dir(directory)

    doc = Document()  # 创建Word文档

    write_to_word(doc, items)  # 写入内容

    name = os.path.basename(directory)
    doc.save(name + '.docx')  # 保存Word文档
    file_path = os.path.join(directory, name + '.docx')

    print(f"Word文档已创建,路径：{file_path}")
